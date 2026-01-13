"""
DOCX Content Extractor
Extracts all text content from a DOCX file including paragraphs, tables, headers, and footers.
"""
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import List, Dict, Any, Tuple


class ContentBlock:
    """Represents a block of content with its type and location."""
    def __init__(self, block_type: str, content: str, index: int, metadata: Dict[str, Any] = None):
        self.block_type = block_type  # 'paragraph', 'table_cell', 'header', 'footer'
        self.content = content
        self.index = index
        self.metadata = metadata or {}
        self.cleaned_content: str = None
    
    def to_dict(self) -> Dict:
        return {
            "type": self.block_type,
            "index": self.index,
            "original": self.content,
            "cleaned": self.cleaned_content,
            "metadata": self.metadata
        }


class DocxExtractor:
    """Extracts all text content from a DOCX file."""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = Document(file_path)
        self.blocks: List[ContentBlock] = []
    
    def extract_all(self) -> List[ContentBlock]:
        """Extract all text content from the document."""
        self.blocks = []
        index = 0
        
        # Extract main body content
        for element in self.document.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = Paragraph(element, self.document)
                text = para.text.strip()
                if text:
                    self.blocks.append(ContentBlock(
                        block_type="paragraph",
                        content=text,
                        index=index,
                        metadata={"style": para.style.name if para.style else None}
                    ))
                    index += 1
                    
            elif element.tag.endswith('tbl'):  # Table
                table = Table(element, self.document)
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            self.blocks.append(ContentBlock(
                                block_type="table_cell",
                                content=cell_text,
                                index=index,
                                metadata={
                                    "row": row_idx,
                                    "col": col_idx,
                                    "table_index": len([b for b in self.blocks if b.block_type == "table_cell"])
                                }
                            ))
                            index += 1
        
        # Extract headers
        for section in self.document.sections:
            if section.header:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        self.blocks.append(ContentBlock(
                            block_type="header",
                            content=text,
                            index=index
                        ))
                        index += 1
            
            # Extract footers
            if section.footer:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        self.blocks.append(ContentBlock(
                            block_type="footer",
                            content=text,
                            index=index
                        ))
                        index += 1
        
        return self.blocks
    
    def get_full_text(self) -> str:
        """Get all text as a single string for context."""
        if not self.blocks:
            self.extract_all()
        return "\n".join([b.content for b in self.blocks])
    
    def get_blocks_for_cleaning(self) -> List[Dict]:
        """Get blocks in a format suitable for LLM processing."""
        if not self.blocks:
            self.extract_all()
        return [{"id": b.index, "text": b.content} for b in self.blocks]
