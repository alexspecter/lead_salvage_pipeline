"""
Generate a test DOCX file with messy content for testing.
"""
from docx import Document
from docx.shared import Inches
import os

def generate_test_docx(output_path: str = "input/test_messy.docx"):
    """Create a DOCX with intentionally messy content."""
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    
    doc = Document()
    
    # Title with typo
    doc.add_heading('Employye Contact Listt', 0)
    
    # Messy paragraphs
    doc.add_paragraph('This documnet containss important contact infromation for our team memberss.')
    doc.add_paragraph('')
    doc.add_paragraph('  JOHN   SMTIH  - Senior Devloper at   Acme Corp  ')
    doc.add_paragraph('Emal: john.smth@acme.com   Phone: (555) 123-4567')
    doc.add_paragraph('')
    doc.add_paragraph('JANE DOE   -   product manager')
    doc.add_paragraph('email:    JANE.DOE@ACME.COM    phone: 555.987.6543')
    doc.add_paragraph('')
    doc.add_paragraph('bob willams - DATA   analyst')
    doc.add_paragraph('E-mail: bob_williams @ acme . com | Ph: 555 555 5555')
    
    # Add a table with messy data
    doc.add_heading('Team Roster', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = 'NAMEE'
    hdr[1].text = 'EMIAL'
    hdr[2].text = 'PHOEN'
    
    # Data rows
    row1 = table.rows[1].cells
    row1[0].text = '  alice   johnson  '
    row1[1].text = 'ALICE.JOHNSON@ACME.COM'
    row1[2].text = '(555)111-2222'
    
    row2 = table.rows[2].cells
    row2[0].text = 'CHARLIE BROWN'
    row2[1].text = 'charlie [at] acme [dot] com'
    row2[2].text = 'call me at 555-333-4444'
    
    row3 = table.rows[3].cells
    row3[0].text = 'david lee'
    row3[1].text = 'david.lee@acme.com'
    row3[2].text = '555.444.5555'
    
    # Footer paragraph
    doc.add_paragraph('')
    doc.add_paragraph('Lastt updatted: Januaryy 13th, 2026    by HR Departmnet')
    
    doc.save(output_path)
    print(f"Generated test DOCX: {output_path}")
    return output_path

if __name__ == "__main__":
    generate_test_docx()
