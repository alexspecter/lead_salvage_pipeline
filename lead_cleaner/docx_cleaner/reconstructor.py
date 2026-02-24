"""
DOCX Reconstructor
Rebuilds a DOCX file with cleaned content while preserving formatting.
"""

from docx import Document
from typing import List, Dict
import os


class DocxReconstructor:
    """Reconstructs a DOCX file with cleaned content."""

    def __init__(self, original_path: str, cleaned_blocks: List[Dict]):
        """
        Args:
            original_path: Path to the original DOCX file
            cleaned_blocks: List of {id, original, cleaned} dicts
        """
        self.original_path = original_path
        self.document = Document(original_path)

        # Build lookup: original_text -> cleaned_text
        self.replacements = {}
        for block in cleaned_blocks:
            if block.get("cleaned") and block.get("original"):
                self.replacements[block["original"]] = block["cleaned"]

    def reconstruct(self, output_path: str):
        """Rebuild the document with cleaned content."""

        # Process main body paragraphs
        for para in self.document.paragraphs:
            original_text = para.text.strip()
            if original_text in self.replacements:
                self._replace_paragraph_text(para, self.replacements[original_text])

        # Process tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text.strip()
                        if original_text in self.replacements:
                            self._replace_paragraph_text(
                                para, self.replacements[original_text]
                            )

        # Process headers and footers
        for section in self.document.sections:
            if section.header:
                for para in section.header.paragraphs:
                    original_text = para.text.strip()
                    if original_text in self.replacements:
                        self._replace_paragraph_text(
                            para, self.replacements[original_text]
                        )

            if section.footer:
                for para in section.footer.paragraphs:
                    original_text = para.text.strip()
                    if original_text in self.replacements:
                        self._replace_paragraph_text(
                            para, self.replacements[original_text]
                        )

        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        # Save the modified document
        self.document.save(output_path)
        return output_path

    def _replace_paragraph_text(self, paragraph, new_text: str):
        """
        Replace text in a paragraph while trying to preserve formatting.
        This is a simplified approach - complex formatting may be lost.
        """
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ""

        # Add new text to first run or create one
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
